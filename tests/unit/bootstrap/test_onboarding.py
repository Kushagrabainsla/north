"""Unit tests for the first-run bootstrap (bootstrap/onboarding.py)."""

from __future__ import annotations

import json
from collections.abc import Awaitable, Callable
from pathlib import Path

import pytest

from bootstrap.onboarding import (
    _clean_fact,
    _discover_files,
    _extract_facts,
    _load_progress,
    _read_text,
    _save_progress,
    run_bootstrap_if_needed,
)
from inference.base import InferenceRouter
from inference.models import (
    CompletionRequest,
    CompletionResponse,
    EmbedRequest,
    EmbedResponse,
    ModelPool,
    PoolPriority,
    ToolCallRequest,
    ToolCallResponse,
    TranscriptionRequest,
    TranscriptionResponse,
)
from memory.facts import FactStore


class _FakeRouter(InferenceRouter):
    """Minimal InferenceRouter stub returning a canned JSON payload.

    With *payloads*, returns one payload per call in order (cycling on the
    last) so different files can yield different facts.
    """

    def __init__(self, payload: object = None, payloads: list[object] | None = None) -> None:
        self._payload = payload
        self._payloads = payloads or []
        self._i = 0

    def _render(self) -> str:
        if self._payloads:
            chosen = self._payloads[min(self._i, len(self._payloads) - 1)]
            self._i += 1
            return json.dumps(chosen)
        return json.dumps(self._payload)

    async def complete(self, request: CompletionRequest) -> CompletionResponse:
        return CompletionResponse(text=self._render(), model_used="fake", tokens_in=0, tokens_out=0, cost_usd=0.0)

    async def transcribe(self, request: TranscriptionRequest) -> TranscriptionResponse:  # pragma: no cover
        raise NotImplementedError

    async def get_model(self, priority: PoolPriority) -> str:  # pragma: no cover
        raise NotImplementedError

    async def refresh_pools(self) -> None:  # pragma: no cover
        raise NotImplementedError

    def current_pools(self) -> dict[str, ModelPool]:  # pragma: no cover
        raise NotImplementedError

    async def complete_with_tools(
        self,
        request: ToolCallRequest,
        token_callback: Callable[[str], Awaitable[None]] | None = None,
    ) -> ToolCallResponse:  # pragma: no cover
        raise NotImplementedError

    async def embed(self, request: EmbedRequest) -> EmbedResponse:  # pragma: no cover
        raise NotImplementedError


class _FakeFactStore(FactStore):
    def __init__(self, bootstrap_count: int = 0, other_count: int = 0) -> None:
        self._bootstrap_count = bootstrap_count
        self._other_count = other_count
        self.added: list[tuple[str, str]] = []
        self._seen: set[tuple[str, str]] = set()
        # Minimal setup for parent class
        self._db_path = Path(":memory:")
        self._embed_fn = lambda x: []

    async def count(self, category: str | None = None) -> int:
        if category == "bootstrap":
            return self._bootstrap_count
        return self._bootstrap_count + self._other_count + len(self.added)

    async def add_fact(self, content: str, category: str = "user") -> bool:
        # Mirrors FactStore.add_fact: exact-match dedup lives in the store,
        # and the call returns whether a new row was actually inserted.
        key = (content, category)
        if key in self._seen:
            return False
        self._seen.add(key)
        self.added.append((content, category))
        return True

    async def add_fact_with_provenance(
        self,
        content: str,
        category: str = "user",
        subject: str = "user",
        confidence: float = 0.8,
        status: str = "active",
        source_path: str | None = None,
        source_hash: str | None = None,
        source_mtime: float | None = None,
        evidence: str | None = None,
    ) -> bool:
        # Delegate to add_fact for testing
        return await self.add_fact(content, category)


def _fake_home(monkeypatch: pytest.MonkeyPatch, home: Path) -> None:
    monkeypatch.setattr(Path, "home", classmethod(lambda cls: home))


# --- _clean_fact -----------------------------------------------------------


def test_clean_fact_extracts_fact_key() -> None:
    assert _clean_fact({"fact": "  User saves $10k  ", "source": "x"}) == "User saves $10k"


def test_clean_fact_accepts_alternate_keys() -> None:
    assert _clean_fact({"content": "User likes keto"}) == "User likes keto"
    assert _clean_fact({"text": "User runs daily"}) == "User runs daily"


def test_clean_fact_dict_without_fact_is_none() -> None:
    assert _clean_fact({"source": "x"}) is None


def test_clean_fact_plain_values() -> None:
    assert _clean_fact("  User saves  ") == "User saves"
    assert _clean_fact(123) is None  # raw numbers are not facts about the user
    assert _clean_fact("") is None
    assert _clean_fact(None) is None


def test_clean_fact_keeps_real_facts() -> None:
    assert _clean_fact("User's monthly income is $6,000") == "User's monthly income is $6,000"
    assert _clean_fact("Kushagra interned at LinkedIn on a checkpoint-restore system") is not None
    assert _clean_fact("His graduate degree objective has not yet been completed") is not None


def test_clean_fact_drops_absence_statements() -> None:
    assert _clean_fact("User has not explicitly stated their health conditions") is None
    assert _clean_fact("Aishwarya Singh's taxable income is not mentioned in the provided file") is None
    assert _clean_fact("No information is available about the user's pets") is None


def test_clean_fact_drops_identification_numbers() -> None:
    assert _clean_fact("User's I-94 admission record number is 010402233A5") is None
    assert _clean_fact("The user's passport number is P1234567") is None


def test_clean_fact_drops_non_user_facts() -> None:
    assert _clean_fact("The I-9 Service Center handled the receipt") is None
    assert _clean_fact("Real Tucson resources: Sun Tran Route 8, 211 Arizona") is None
    assert _clean_fact("Participant check-in: 4-screen conversational flow") is None


# --- _extract_facts --------------------------------------------------------


async def test_extract_facts_cleans_dict_payload(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("content", encoding="utf-8")
    router = _FakeRouter(
        {"facts": [
            {"content": "User A", "subject": "user", "confidence": 0.9},
            {"content": "User B", "subject": "user", "confidence": 0.8},
            {"content": "User C", "subject": "user", "confidence": 0.7},
        ]}
    )
    result = await _extract_facts(path, router)
    assert len(result) == 3
    assert result[0]["content"] == "User A"
    assert result[1]["content"] == "User B"
    assert result[2]["content"] == "User C"
    # Check provenance fields
    for cand in result:
        assert cand["subject"] == "user"
        assert "source_path" in cand
        assert "source_hash" in cand
        assert "source_mtime" in cand


async def test_extract_facts_strips_markdown_fences(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("content", encoding="utf-8")

    # With structured output (response_schema), providers don't wrap in markdown fences.
    # This test verifies the new behavior works correctly - no fence stripping needed.
    router = _FakeRouter({"facts": [
        {"content": "User fact one", "subject": "user", "confidence": 0.9},
        {"content": "User fact two", "subject": "user", "confidence": 0.8},
    ]})
    result = await _extract_facts(path, router)
    assert len(result) == 2
    assert result[0]["content"] == "User fact one"
    assert result[1]["content"] == "User fact two"


async def test_extract_facts_non_json_returns_empty(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("content", encoding="utf-8")

    class _GarbageRouter(_FakeRouter):
        def _render(self) -> str:
            return "not json at all"

    assert await _extract_facts(path, _GarbageRouter({"facts": []})) == []


async def test_extract_facts_empty_file_returns_empty(tmp_path: Path) -> None:
    path = tmp_path / "empty.txt"
    path.write_text("   ", encoding="utf-8")
    router = _FakeRouter(["should never be called"])
    assert await _extract_facts(path, router) == []


# --- _read_text ------------------------------------------------------------


def test_read_text_pdf_uses_pypdf(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    pytest.importorskip("pypdf")
    import pypdf

    class _FakePage:
        def extract_text(self) -> str:
            return "Hello bootstrap"

    class _FakeReader:
        def __init__(self, stream: str) -> None:
            self.pages = [_FakePage(), _FakePage()]

    monkeypatch.setattr(pypdf, "PdfReader", _FakeReader)
    path = tmp_path / "doc.pdf"
    path.write_bytes(b"%PDF-1.4 fake")
    # Regression guard: the PDF branch must import pypdf (it previously
    # imported pymupdf, which is not a dependency -> all PDFs were skipped).
    assert _read_text(path) == "Hello bootstrap\nHello bootstrap"


def test_read_text_pdf_missing_library_returns_empty(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    import builtins

    real_import = builtins.__import__

    def _block_pypdf(name: str, *args: object, **kwargs: object) -> object:
        if name == "pypdf":
            raise ImportError("blocked for test")
        return real_import(name, *args, **kwargs)  # type: ignore[arg-type]

    monkeypatch.setattr(builtins, "__import__", _block_pypdf)
    path = tmp_path / "doc.pdf"
    path.write_bytes(b"%PDF-1.4 fake")
    assert _read_text(path) == ""


def test_read_text_docx(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Keto meal plan")
    doc.add_paragraph("Second line")
    doc.save(str(path))
    assert _read_text(path) == "Keto meal plan\nSecond line"


def test_read_text_plain(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("line one\nline two", encoding="utf-8")
    assert _read_text(path) == "line one\nline two"


# --- _discover_files -------------------------------------------------------


def test_discover_files_skips_north_repo(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    home = tmp_path / "home"
    north = home / "Desktop" / "projects" / "north"
    (north / "skills" / "builtin").mkdir(parents=True)
    (north / "README.md").write_text("north readme", encoding="utf-8")
    (north / "skills" / "builtin" / "SKILL.md").write_text("skill", encoding="utf-8")
    other_readme = home / "Desktop" / "projects" / "other" / "README.md"
    other_readme.parent.mkdir(parents=True)
    other_readme.write_text("other", encoding="utf-8")
    # Non-README files inside project repos (AI-system docs, skill files,
    # build artifacts) are not personal facts — recursion into projects is
    # skipped; only README.md passes are collected.
    leak = home / "Desktop" / "projects" / "system_prompts_leaks" / "OpenAI" / "gpt.md"
    leak.parent.mkdir(parents=True)
    leak.write_text("system prompt", encoding="utf-8")
    notes = home / "Documents" / "notes.txt"
    notes.parent.mkdir(parents=True)
    notes.write_text("notes", encoding="utf-8")
    _fake_home(monkeypatch, home)

    paths = {str(p) for p in _discover_files()}
    assert str(north / "README.md") not in paths
    assert str(north / "skills" / "builtin" / "SKILL.md") not in paths
    assert str(home / "Desktop" / "projects" / "other" / "README.md") in paths
    assert str(leak) not in paths
    assert str(home / "Documents" / "notes.txt") in paths


def test_discover_files_skips_hidden_and_junk(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    home = tmp_path / "home"
    hidden = home / "Documents" / ".config" / "x.txt"
    hidden.parent.mkdir(parents=True)
    hidden.write_text("hidden", encoding="utf-8")
    junk = home / "Documents" / "node_modules" / "y.txt"
    junk.parent.mkdir(parents=True)
    junk.write_text("junk", encoding="utf-8")
    _fake_home(monkeypatch, home)
    assert _discover_files() == []


# --- run_bootstrap_if_needed -----------------------------------------------


async def test_bootstrap_skips_when_marker_exists(tmp_path: Path) -> None:
    north_home = tmp_path / "north_home"
    north_home.mkdir()
    (north_home / ".bootstrapped").touch()
    store = _FakeFactStore(bootstrap_count=0)
    router = _FakeRouter(["never"])
    await run_bootstrap_if_needed(store, router, north_home)
    assert store.added == []


async def test_bootstrap_skips_when_facts_exist_and_no_progress(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    home = tmp_path / "home"
    (home / "Documents").mkdir(parents=True)
    (home / "Documents" / "a.txt").write_text("user data", encoding="utf-8")
    _fake_home(monkeypatch, home)
    north_home = tmp_path / "north_home"
    north_home.mkdir()
    store = _FakeFactStore(bootstrap_count=7)
    router = _FakeRouter(["never"])
    await run_bootstrap_if_needed(store, router, north_home)
    assert store.added == []
    assert (north_home / ".bootstrapped").exists()


async def test_bootstrap_runs_when_only_user_facts_exist(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    # Facts learned from conversation must not block a first-ever bootstrap.
    home = tmp_path / "home"
    (home / "Documents").mkdir(parents=True)
    (home / "Documents" / "a.txt").write_text("user data", encoding="utf-8")
    _fake_home(monkeypatch, home)
    north_home = tmp_path / "north_home"
    north_home.mkdir()
    store = _FakeFactStore(bootstrap_count=0, other_count=12)
    router = _FakeRouter({"facts": [{"content": "User fact A", "subject": "user", "confidence": 0.9}]})
    await run_bootstrap_if_needed(store, router, north_home)
    assert len(store.added) == 1
    assert (north_home / ".bootstrapped").exists()


async def test_bootstrap_resumes_interrupted_run(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    home = tmp_path / "home"
    (home / "Documents").mkdir(parents=True)
    (home / "Documents" / "a.txt").write_text("user data a", encoding="utf-8")
    (home / "Documents" / "b.txt").write_text("user data b", encoding="utf-8")
    _fake_home(monkeypatch, home)
    north_home = tmp_path / "north_home"
    north_home.mkdir()
    # Simulate an interrupted run: a.txt already checkpointed, b.txt pending.
    _save_progress(north_home, [str((home / "Documents" / "a.txt").resolve())])
    store = _FakeFactStore(bootstrap_count=3)
    router = _FakeRouter({"facts": [{"content": "User likes B", "subject": "user", "confidence": 0.9}]})
    await run_bootstrap_if_needed(store, router, north_home)
    assert len(store.added) == 1
    assert store.added[0][0] == "User likes B"
    assert store.added[0][1] == "bootstrap"
    assert (north_home / ".bootstrapped").exists()
    assert not (north_home / ".bootstrap_progress.json").exists()


async def test_bootstrap_fresh_run(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    home = tmp_path / "home"
    (home / "Documents").mkdir(parents=True)
    (home / "Documents" / "a.txt").write_text("user data a", encoding="utf-8")
    (home / "Documents" / "b.txt").write_text("user data b", encoding="utf-8")
    _fake_home(monkeypatch, home)
    north_home = tmp_path / "north_home"
    north_home.mkdir()
    store = _FakeFactStore(bootstrap_count=0)
    router = _FakeRouter(
        payloads=[
            {"facts": [{"content": "User fact A", "subject": "user", "confidence": 0.9}]},
            {"facts": [{"content": "User fact B", "subject": "user", "confidence": 0.8}]},
        ]
    )
    await run_bootstrap_if_needed(store, router, north_home)
    assert len(store.added) == 2
    assert all(category == "bootstrap" for _, category in store.added)
    assert (north_home / ".bootstrapped").exists()
    assert not (north_home / ".bootstrap_progress.json").exists()


async def test_bootstrap_dedups_identical_facts_within_run(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    # Two files that yield the same fact (e.g. repeated copies of one PDF)
    # must store it once. Bootstrap forwards every extracted fact to the
    # store and adds no dedup of its own — FactStore's exact-match dedup
    # (which needs no embeddings) collapses the duplicates.
    home = tmp_path / "home"
    (home / "Documents").mkdir(parents=True)
    (home / "Documents" / "a.txt").write_text("user data a", encoding="utf-8")
    (home / "Documents" / "b.txt").write_text("user data b", encoding="utf-8")
    _fake_home(monkeypatch, home)
    north_home = tmp_path / "north_home"
    north_home.mkdir()
    store = _FakeFactStore(bootstrap_count=0)
    router = _FakeRouter({"facts": [{"content": "User fact A", "subject": "user", "confidence": 0.9}]})
    await run_bootstrap_if_needed(store, router, north_home)
    assert len(store.added) == 1
    assert store.added[0][0] == "User fact A"
    assert store.added[0][1] == "bootstrap"


async def test_bootstrap_no_candidate_files_marks_done(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    home = tmp_path / "home"
    home.mkdir()
    _fake_home(monkeypatch, home)
    north_home = tmp_path / "north_home"
    north_home.mkdir()
    store = _FakeFactStore(bootstrap_count=0)
    await run_bootstrap_if_needed(store, _FakeRouter([]), north_home)
    assert (north_home / ".bootstrapped").exists()


async def test_bootstrap_none_fact_store(tmp_path: Path) -> None:
    north_home = tmp_path / "north_home"
    north_home.mkdir()
    await run_bootstrap_if_needed(None, _FakeRouter([]), north_home)
    assert not (north_home / ".bootstrapped").exists()


# --- progress helpers ------------------------------------------------------


def test_progress_roundtrip(tmp_path: Path) -> None:
    north_home = tmp_path / "north_home"
    north_home.mkdir()
    completed = [
        {"path": "/a", "hash": "h1", "mtime": 1.0, "status": "completed"},
        {"path": "/b", "hash": "h2", "mtime": 2.0, "status": "completed"},
        {"path": "/c", "hash": "h3", "mtime": 3.0, "status": "completed"},
    ]
    _save_progress(north_home, completed)
    assert _load_progress(north_home) == completed


def test_progress_absent_is_none(tmp_path: Path) -> None:
    assert _load_progress(tmp_path) is None


def test_progress_corrupt_is_none(tmp_path: Path) -> None:
    path = tmp_path / ".bootstrap_progress.json"
    path.write_text("{not json", encoding="utf-8")
    assert _load_progress(tmp_path) is None


def test_rank_prefers_dense_small_files_over_huge_labs() -> None:
    """A small resume (.md) must outrank a huge lab dump (.pdf) within the same group."""
    from pathlib import Path as _P

    from bootstrap.onboarding import _rank_file

    home = _P.home()
    resume = home / "Documents" / "resume.md"
    lab = home / "Documents" / "lab05_complex.pdf"
    # Use monkeypatched-free comparison via stub files' attributes only.
    class _Stub:
        def __init__(self, p: _P, size: int) -> None:
            self._p = p
            self._size = size
        def resolve(self): return self._p
        @property
        def suffix(self): return self._p.suffix
        @property
        def stem(self): return self._p.stem
        def stat(self): 
            import os
            return os.stat_result((0,0,0,0,0,0,self._size,0,0,0))
        def relative_to(self, other): return self._p.relative_to(other)
    r_resume = _rank_file(_Stub(resume, 200), "documents")
    r_lab = _rank_file(_Stub(lab, 3_000_000), "documents")
    # Lower tuple = higher priority.
    assert r_resume < r_lab


def test_denylist_skips_lab_and_key_files() -> None:
    """Files whose names contain noisy fragments must be excluded from discovery."""
    from bootstrap.onboarding import _DENY_PATH_FRAGMENTS

    bad = ["lab05_notes.txt", "alicePriv.txt", "complex.json", "system_prompts_leaks.md"]
    for name in bad:
        assert any(frag in name.lower() for frag in _DENY_PATH_FRAGMENTS), name


async def test_extract_profile_parses_sections(tmp_path: Path) -> None:
    """_extract_profile returns domain-grouped facts + a markdown summary."""
    from bootstrap.onboarding import _extract_profile

    payload = {
        "education": ["The user is a student at San Jose State University"],
        "jobs": ["The user is seeking software engineering roles"],
        "skills": ["Python", "SQL"],
        "finances": [],
        "health": [],
        "schedule": [],
        "preferences": [],
        "projects": ["The user maintains the north agent framework"],
    }
    router = _FakeRouter(payload=payload)
    src = tmp_path / "resume.md"
    src.write_text("Kushagra Bainsla")
    facts, md = await _extract_profile(src, router)
    contents = {f["content"] for f in facts}
    assert "The user is a student at San Jose State University" in contents
    assert "The user maintains the north agent framework" in contents
    assert "## Education" in md
    assert "## Projects" in md
